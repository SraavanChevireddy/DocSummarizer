import os
from openai import OpenAI
from flask import Flask, request, jsonify
from transformers import pipeline, AutoTokenizer, AutoModelForTokenClassification
import torch
import PyPDF2
from docx import Document

app = Flask(__name__)

# OpenAI API Configuration
openai_client = OpenAI(api_key="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")

# Hugging Face Summarizer Configuration
device = "mps" if torch.backends.mps.is_available() else "cpu"
print(f"Using device for Hugging Face models: {device}")

summarizer = pipeline(
    "summarization",
    model="facebook/bart-large-cnn",
    device=0 if device == "mps" else -1
)

# Function to Extract Text from PDF Files
def extract_text_from_pdf(file_path):
    text = ""
    try:
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                text += page.extract_text()
    except Exception as e:
        print(f"Error reading PDF file: {e}")
    return text

# Function to Extract Text from DOCX Files
def extract_text_from_docx(file_path):
    text = ""
    try:
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
    except Exception as e:
        print(f"Error reading DOCX file: {e}")
    return text

# Helper Function to Chunk Text for Hugging Face Summarizer
def chunk_text(text, max_chunk_size=800):
    sentences = text.split(". ")
    chunks = []
    chunk = ""

    for sentence in sentences:
        if len(chunk) + len(sentence) <= max_chunk_size:
            chunk += sentence + ". "
        else:
            chunks.append(chunk.strip())
            chunk = sentence + ". "

    if chunk:
        chunks.append(chunk.strip())
    return chunks

# Hugging Face Summarization Process
def summarize_using_huggingface(text, max_length=200, min_length=100):
    try:
        summary = summarizer(
            text,
            max_length=max_length,
            min_length=min_length,
            truncation=True
        )
        return summary[0]['summary_text']
    except Exception as e:
        print(f"Error summarizing with Hugging Face: {e}")
        return ""

def summarize_document_huggingface(text):
    chunks = chunk_text(text, max_chunk_size=800)
    combined_summary = ""

    print(f"Text is split into {len(chunks)} chunks for summarization...")
    for i, chunk in enumerate(chunks):
        print(f"Summarizing chunk {i + 1}/{len(chunks)}...")
        combined_summary += summarize_using_huggingface(chunk) + " "

    print("\nGenerating final overall summary...")
    final_summary = summarize_using_huggingface(combined_summary.strip())
    return final_summary

# OpenAI Summarization Process
def summarize_using_openai(text, max_length=150):
    try:
        response = openai_client.chat.completions.create(
            model="gpt-4o-mini",  # Replace with "gpt-3.5-turbo" or other models based on your preference
            messages=[
                {"role": "system", "content": "You are a helpful assistant that summarizes text."},
                {"role": "user", "content": f"Please summarize the following text:\n\n{text}"}
            ],
            temperature=0.7,
            max_tokens=max_length
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        print(f"Error using OpenAI API for summarization: {e}")
        return ""

# Flask Route for OpenAI Summarization
@app.route('/summarize/premium', methods=['POST'])
def summarize_openai():
    if 'file' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files['file']
    file_ext = os.path.splitext(file.filename)[1].lower()

    text = ""
    try:
        if file_ext == ".pdf":
            text = extract_text_from_pdf(file)
        elif file_ext in [".doc", ".docx"]:
            text = extract_text_from_docx(file)
        else:
            return jsonify({"error": "Unsupported file format"}), 400
    except Exception as e:
        return jsonify({"error": f"Error processing file: {e}"}), 500

    if not text.strip():
        return jsonify({"error": "No text extracted from the file"}), 400

    summary = summarize_using_openai(text, max_length=200)
    return jsonify({"summary": summary})

# Flask Route for Hugging Face Summarization
@app.route('/summarize/basic', methods=['POST'])
def summarize_huggingface():
    print("Route /summarize/basic was accessed")
    if 'file' not in request.files:
        print("No file uploaded")
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files['file']
    file_ext = os.path.splitext(file.filename)[1].lower()

    text = ""
    try:
        if file_ext == ".pdf":
            text = extract_text_from_pdf(file)
        elif file_ext in [".doc", ".docx"]:
            text = extract_text_from_docx(file)
        else:
            return jsonify({"error": "Unsupported file format"}), 400
    except Exception as e:
        return jsonify({"error": f"Error processing file: {e}"}), 500

    if not text.strip():
        return jsonify({"error": "No text extracted from the file"}), 400

    summary = summarize_document_huggingface(text)
    return jsonify({"summary": summary})

# BioBERT Model Setup: Load BioBERT for named entity extraction
model_name = "dmis-lab/biobert-base-cased-v1.1"  # BioBERT model
tokenizer = AutoTokenizer.from_pretrained(model_name)
model = AutoModelForTokenClassification.from_pretrained(model_name)

ner_pipeline = pipeline(
    "ner",
    model=model,
    tokenizer=tokenizer,
    device=0 if torch.backends.mps.is_available() else -1,
    aggregation_strategy="simple"
)

print("BioBERT model loaded successfully!")

# Function to Extract Text from PDF Files
def extract_text_from_pdf(file_path):
    text = ""
    try:
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                text += page.extract_text()
    except Exception as e:
        print(f"Error reading PDF file: {e}")
    return text

# Function to Extract Text from DOCX Files
def extract_text_from_docx(file_path):
    text = ""
    try:
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
    except Exception as e:
        print(f"Error reading DOCX file: {e}")
    return text

# Function to Use BioBERT for Medical Keyword Extraction
def extract_medical_keywords(text):
    try:
        ner_results = ner_pipeline(text)
        medical_keywords = {
            "keywords": [entity['word'] for entity in ner_results if entity['entity'] == 'B-MEDICAL' or 'MED']
        }
        return medical_keywords
    except Exception as e:
        print(f"Error extracting keywords with BioBERT: {e}")
        return {"error": str(e)}

# Flask Route for Medical Keyword Extraction using BioBERT
@app.route('/extract-medical-keywords', methods=['POST'])
def extract_keywords():
    if 'file' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files['file']
    file_ext = os.path.splitext(file.filename)[1].lower()

    text = ""
    try:
        if file_ext == ".pdf":
            text = extract_text_from_pdf(file)
        elif file_ext in [".doc", ".docx"]:
            text = extract_text_from_docx(file)
        else:
            return jsonify({"error": "Unsupported file format"}), 400
    except Exception as e:
        return jsonify({"error": f"Error processing file: {e}"}), 500

    if not text.strip():
        return jsonify({"error": "No text extracted from the file"}), 400

    # Extract medical-related keywords using BioBERT
    medical_keywords = extract_medical_keywords(text)
    return jsonify(medical_keywords)

@app.route('/ping', methods=['GET'])
def ping():
    return jsonify({"message": "Connection successful!"}), 200


@app.route('/routes', methods=['GET'])
def list_routes():
    output = []
    for rule in app.url_map.iter_rules():
        output.append({
            "endpoint": rule.endpoint,
            "methods": list(rule.methods),
            "route": str(rule)
        })
    return jsonify(output)


# Run the Flask App
if __name__ == '__main__':
    app.run(debug=True)
